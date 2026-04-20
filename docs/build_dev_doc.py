"""docs/개발문서.docx 생성 스크립트."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "개발문서.docx"

KOR_FONT = "맑은 고딕"
CODE_FONT = "Consolas"


def set_korean_font(run, size_pt: float, bold: bool = False, color=None) -> None:
    run.font.name = KOR_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)


def set_code_font(run, size_pt: float = 9) -> None:
    run.font.name = CODE_FONT
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), CODE_FONT)
    r_fonts.set(qn("w:ascii"), CODE_FONT)
    r_fonts.set(qn("w:hAnsi"), CODE_FONT)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    align=None,
    space_after: float = 6,
    color=None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_korean_font(run, size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    if level == 0:
        size = 22
    elif level == 1:
        size = 16
    elif level == 2:
        size = 13
    else:
        size = 12
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_korean_font(run, size, bold=True, color=RGBColor(0x1F, 0x2D, 0x3D))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        if p.runs:
            run = p.runs[0]
            run.text = item
        else:
            run = p.add_run(item)
        set_korean_font(run, 11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        if p.runs:
            run = p.runs[0]
            run.text = item
        else:
            run = p.add_run(item)
        set_korean_font(run, 11)


def add_code_block(doc: Document, code: str) -> None:
    """회색 배경 + Consolas 9pt 코드 블록을 1x1 표로 구현."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(15.8)
    cell = table.cell(0, 0)
    cell.width = Cm(15.8)
    _set_cell_shading(cell, "F4F5F7")
    cell.text = ""
    lines = code.splitlines() or [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(line if line else " ")
        set_code_font(run, 9)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_korean_font(run, 10.5, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_korean_font(run, 10.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(
        doc,
        "Cost Compass 개발문서",
        size=26,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x10, 0x2A, 0x43),
        space_after=4,
    )
    add_paragraph(
        doc,
        "아키텍처와 바이브코딩 프로세스 (Interface Hub 재활용)",
        size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(
        doc,
        "노아에이티에스 2025년 상반기 연구소 인력 채용 포트폴리오 (주제 2)",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x44, 0x55, 0x66),
        space_after=80,
    )
    add_paragraph(
        doc,
        "작성자: 홍유진 (GitHub: dbwls99706)",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "작성일: 2026-04-20",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "라이브 데모: https://cost-compass.vercel.app",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "저장소: https://github.com/dbwls99706/cost-compass",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "자매 프로젝트: https://github.com/dbwls99706/interface-hub",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_heading(doc, "목차", level=0)
    items = [
        "1. 문서 목적과 범위",
        "2. 시스템 아키텍처",
        "3. 데이터 모델",
        "4. 집계 쿼리 설계",
        "5. 프런트엔드 데이터 흐름",
        "6. URL 필터와 드릴다운",
        "7. 로컬 SQLite와 Turso 전환 전략",
        "8. 바이브코딩 프로세스",
        "9. 검증 전략",
        "10. 트러블슈팅 사례",
        "11. 부록",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, space_after=4)
    doc.add_page_break()


def section_scope(doc: Document) -> None:
    add_heading(doc, "1. 문서 목적과 범위", level=1)
    add_paragraph(
        doc,
        "이 문서는 Cost Compass의 설계 결정과 개발 프로세스를 심사위원이 재현 가능한 수준으로 기술한다. "
        "제품의 기능 명세는 기획서(docs/기획서.docx)를 참조한다. Cost Compass는 자매 프로젝트인 "
        "Interface Hub에서 이미 검증된 기술 스택, 배포 파이프라인, 바이브코딩 프로세스를 재활용해 "
        "반나절 만에 완성한 두 번째 산출물이다.",
    )
    add_bullets(
        doc,
        [
            "대상 독자: 채용 심사위원, 신규 합류 개발자.",
            "다루는 범위: 시스템 아키텍처, 데이터 모델, 집계 쿼리, URL 드릴다운, 바이브코딩 프로세스, 트러블슈팅.",
            "다루지 않는 범위: 사용자 가이드, UI 카피 리뷰, 화면별 인터랙션 세부 스펙.",
        ],
    )


def section_architecture(doc: Document) -> None:
    add_heading(doc, "2. 시스템 아키텍처", level=1)

    add_heading(doc, "2.1 고수준 구성", level=2)
    add_paragraph(
        doc,
        "브라우저는 Next.js 서버와 Server Action으로 통신한다. 서버는 Prisma Client를 통해 libSQL "
        "driver adapter로 DB에 접근한다. 로컬 개발에서는 SQLite 파일, 프로덕션에서는 Turso(libSQL) "
        "인스턴스를 사용한다. Interface Hub와 동일한 파이프라인이다.",
    )
    add_code_block(
        doc,
        "Browser\n"
        "   │  fetch / Link navigation\n"
        "   ▼\n"
        "Next.js (Server Component)\n"
        "   │  getDashboardStats / listProjects / getProjectDetail\n"
        "   │  listDivisions / getDivisionDetail\n"
        "   ▼\n"
        "Server Action (lib/actions/cost-query.ts)\n"
        "   │  prisma.{division,project,costItem}.findMany\n"
        "   ▼\n"
        "JS 집계 (groupBy 없이 Map 기반)\n"
        "   │\n"
        "   ▼\n"
        "Client Chart (Recharts) - 표시만 담당",
    )

    add_heading(doc, "2.2 주요 결정 사항", level=2)
    add_table(
        doc,
        headers=["결정 지점", "선택", "이유"],
        rows=[
            [
                "App Router vs Pages Router",
                "App Router",
                "Server Component로 집계 쿼리를 서버에서 완결하고, 클라이언트는 차트 표시만 담당한다.",
            ],
            [
                "CRUD vs 조회 전용",
                "조회 전용",
                "Interface Hub가 '쓰는' 도구였다. Cost Compass는 '읽는' 도구로 포지셔닝한다.",
            ],
            [
                "Prisma groupBy vs JS 집계",
                "JS 집계",
                "SQLite + driver adapter 호환성 확보, 326건 규모에서 CPU/메모리 부담 무시 가능.",
            ],
            [
                "Prisma + libSQL Adapter",
                "Turso 전환 용이",
                "런타임 환경변수 하나로 로컬 SQLite와 Turso를 분기한다.",
            ],
            [
                "필터 상태 관리",
                "URL 쿼리스트링",
                "공유 가능한 링크와 Server Component 기반 구조에 자연스럽게 맞는다.",
            ],
            [
                "차트 라이브러리",
                "Recharts",
                "5종 차트(Line, Pie, Stacked Bar, Horizontal Bar, Table)를 단일 API로 구성.",
            ],
        ],
        col_widths_cm=[4.2, 3.8, 7.8],
    )


def section_data_model(doc: Document) -> None:
    add_heading(doc, "3. 데이터 모델", level=1)

    add_heading(doc, "3.1 엔티티", level=2)
    add_paragraph(doc, "3개 엔티티로 전사 원가 데이터를 정규화한다.")
    add_table(
        doc,
        headers=["엔티티", "주요 필드"],
        rows=[
            [
                "Division",
                "id, name (unique), headcount, createdAt",
            ],
            [
                "Project",
                "id, code (unique, PRJ-001 형식), name, divisionId, status, startedAt, endedAt, description",
            ],
            [
                "CostItem",
                "id, projectId, divisionId, sourceDivisionId?, category, period, standardAmount, actualAmount, isInterUnit, notes",
            ],
        ],
        col_widths_cm=[3.0, 12.8],
    )

    add_heading(doc, "3.2 관계", level=2)
    add_bullets(
        doc,
        [
            "Division 1:N Project (본부 삭제 시 프로젝트 Cascade).",
            "Project 1:N CostItem (프로젝트 삭제 시 원가 항목 Cascade).",
            "Division 1:N CostItem 이중 관계: divisionId는 '부담 본부', sourceDivisionId는 '내부대체 출처 본부'.",
        ],
    )

    add_heading(doc, "3.3 내부대체 표현", level=2)
    add_paragraph(
        doc,
        "CostItem은 두 개의 본부 참조를 갖는다. divisionId는 비용을 최종적으로 부담하는 본부를, "
        "sourceDivisionId는 내부대체가 일어난 경우의 출처 본부를 가리킨다. isInterUnit 불린으로 "
        "내부대체 여부를 명시해 KPI의 '내부대체 비중' 계산을 단순화했다.",
    )

    add_heading(doc, "3.4 주요 인덱스", level=2)
    add_table(
        doc,
        headers=["엔티티", "인덱스", "이유"],
        rows=[
            ["CostItem", "(projectId, period)", "프로젝트 상세의 월별 집계"],
            ["CostItem", "(divisionId, period)", "본부 상세의 월별 집계"],
            ["CostItem", "(category)", "카테고리별 비중 계산"],
            ["CostItem", "(period)", "전사 월별 추이"],
            ["Project", "(divisionId)", "본부별 프로젝트 목록"],
            ["Project", "(status)", "진행 상태 필터"],
        ],
        col_widths_cm=[3.2, 5.6, 7.0],
    )


def section_aggregation(doc: Document) -> None:
    add_heading(doc, "4. 집계 쿼리 설계", level=1)

    add_heading(doc, "4.1 Prisma groupBy를 쓰지 않는 이유", level=2)
    add_bullets(
        doc,
        [
            "SQLite + libSQL driver adapter 환경의 호환성 확보. 일부 집계/JOIN 조합에서 드라이버별 동작 차이가 있어 예측 가능한 JS 집계를 택했다.",
            "326건 규모에서 메모리/CPU 부담이 무시 가능한 수준. 대시보드 1회 요청이 DB 3회 호출로 끝난다.",
            "집계 로직이 타입스크립트 함수 한 곳에 모여 테스트, 디버깅, 수정이 단순해진다.",
            "규모가 커지면 prisma.costItem.aggregate나 raw SQL로 경계 하나만 갈아끼우면 되는 구조로 설계했다 (lib/actions/cost-query.ts).",
        ],
    )

    add_heading(doc, "4.2 Map 기반 집계 패턴", level=2)
    add_paragraph(
        doc,
        "모든 집계 함수는 'Promise.all 병렬 조회 → 필터 → 버킷 누적 → 정렬'의 동일한 골격을 공유한다. "
        "가독성을 해치지 않으면서 카테고리 합산, 월별 합산, 프로젝트별 합산을 모두 같은 패턴으로 처리한다.",
    )
    add_code_block(
        doc,
        "const costBuckets = new Map<string, { actual: number; standard: number }>();\n"
        "for (const item of costItems) {\n"
        "  const bucket = costBuckets.get(item.divisionId)\n"
        "    ?? { actual: 0, standard: 0 };\n"
        "  bucket.actual   += item.actualAmount;\n"
        "  bucket.standard += item.standardAmount;\n"
        "  costBuckets.set(item.divisionId, bucket);\n"
        "}",
    )

    add_heading(doc, "4.3 KPI와 파생값 계산", level=2)
    add_bullets(
        doc,
        [
            "Variance = (실제 - 표준). 비율은 standard가 0일 때 0으로 방어한다.",
            "내부대체 비중 = isInterUnit=true인 항목의 실제 합 / 전체 실제 합.",
            "카테고리 비중 = 카테고리별 실제 합 / 전체 실제 합.",
            "월별 추이는 period 문자열(YYYY-MM)을 사전순으로 정렬하면 시간순이 된다.",
        ],
    )


def section_frontend(doc: Document) -> None:
    add_heading(doc, "5. 프런트엔드 데이터 흐름", level=1)

    add_heading(doc, "5.1 화면별 렌더링 전략", level=2)
    add_table(
        doc,
        headers=["화면", "전략", "이유"],
        rows=[
            ["/", "Server Component", "집계는 SSR에서 완결, 차트만 Client."],
            ["/projects", "Server Component", "목록 테이블과 필터링, 인터랙션은 URL로."],
            ["/projects/[id]", "Server Component", "params Promise await, CSV 버튼만 Client."],
            ["/divisions", "Server Component", "카드 그리드, 각 카드는 Link로 상세로 이동."],
            ["/divisions/[id]", "Server Component", "KPI + 차트 2종 + 프로젝트 목록 모두 SSR."],
        ],
        col_widths_cm=[3.8, 4.2, 7.8],
    )

    add_heading(doc, "5.2 Client 경계", level=2)
    add_bullets(
        doc,
        [
            "차트 컴포넌트(MonthlyTrendChart, CategoryPieChart, DivisionStackedBar, TopProjectsBar)는 Recharts 의존으로 'use client'.",
            "FilterBar는 useRouter로 URL을 갱신해야 해서 Client.",
            "ExportCsvButton은 Server Action을 호출해 Blob 다운로드를 트리거하므로 Client.",
            "나머지 Badge, KpiCard 등 표시 전용 컴포넌트는 Server Component로 유지.",
        ],
    )

    add_heading(doc, "5.3 Recharts 주의 사항", level=2)
    add_bullets(
        doc,
        [
            "ResponsiveContainer에는 minWidth={0} minHeight={0}을 항상 붙인다 (측정 race 방지).",
            "Tooltip의 formatter는 (value, name) 시그니처로 받아 내부에서 Number(), String() 캐스팅한다 (Recharts 3의 타입 엄격화 대응).",
            "onClick 이벤트 파라미터는 정식 타입이 좁아 as unknown을 거쳐 { activePayload?: ... } 좁은 타입으로 좁힌다.",
        ],
    )


def section_drilldown(doc: Document) -> None:
    add_heading(doc, "6. URL 필터와 드릴다운", level=1)

    add_heading(doc, "6.1 URL 쿼리스트링 계약", level=2)
    add_table(
        doc,
        headers=["파라미터", "허용 값", "비고"],
        rows=[
            ["period", "YYYY-MM 또는 all", "기본 all."],
            ["divisionId", "Division.id 또는 all", "기본 all."],
            ["status", "PLANNING / ACTIVE / CLOSED 또는 all", "대소문자 엄격 일치."],
        ],
        col_widths_cm=[3.6, 5.4, 6.8],
    )
    add_paragraph(
        doc,
        "Server Component는 searchParams를 Promise로 받아 await한 뒤 필터 객체로 변환한다. "
        "FilterBar의 onValueChange는 useRouter().push로 쿼리스트링을 갱신해 동일 URL로 직접 공유할 수 있다.",
    )

    add_heading(doc, "6.2 차트 클릭 드릴다운", level=2)
    add_bullets(
        doc,
        [
            "프로젝트 Top 10의 막대 클릭은 /projects/[id]로 이동한다.",
            "본부별 Stacked Bar의 막대 클릭은 /divisions/[id]로 이동한다.",
            "클릭 타겟은 activePayload에서 payload.projectId 또는 payload.divisionId를 꺼내 router.push에 넘긴다.",
            "차트에 cursor: 'pointer' 스타일을 달아 클릭 가능함을 시각적으로 드러낸다.",
        ],
    )


def section_turso(doc: Document) -> None:
    add_heading(doc, "7. 로컬 SQLite와 Turso 전환 전략", level=1)

    add_heading(doc, "7.1 문제", level=2)
    add_bullets(
        doc,
        [
            "Prisma 7 CLI는 libsql:// URL을 직접 이해하지 못한다.",
            "Driver adapter는 런타임에만 주입되고, migrate 커맨드에 공식 경로로 전달할 수 없다.",
            "Vercel 배포 시 Turso에 마이그레이션을 어떻게 적용할지가 별도 과제가 된다.",
        ],
    )

    add_heading(doc, "7.2 해결", level=2)
    add_bullets(
        doc,
        [
            "런타임은 lib/prisma.ts의 resolveLibSqlConfig()에서 file: 과 libsql: 양쪽을 처리한다.",
            "환경변수 우선순위는 TURSO_DATABASE_URL > DATABASE_URL(libsql://) > DATABASE_URL(file:) 순이다.",
            "마이그레이션은 scripts/turso-deploy.ts를 통해 우회한다. libSQL client로 SQL을 직접 실행한다.",
            "Interface Hub에서 이미 검증된 동일 스크립트 구조를 그대로 복제했다.",
        ],
    )

    add_heading(doc, "7.3 Turso 배포 스크립트 원리", level=2)
    add_bullets(
        doc,
        [
            "_applied_migrations 테이블을 만들어 멱등성을 확보한다.",
            "prisma/migrations를 이름 순으로 순회하고, 이미 적용된 마이그레이션은 스킵한다.",
            "주석 라인을 라인 단위로 제거한 뒤 세미콜론 기준으로 statement를 분리해 하나씩 execute한다.",
            "실패 시 롤백 대신 scripts/turso-reset.ts로 적용 기록을 초기화하고 재시도한다.",
        ],
    )
    add_code_block(
        doc,
        "const stripComments = (sql: string): string =>\n"
        "  sql\n"
        "    .split(\"\\n\")\n"
        "    .filter((line) => !line.trim().startsWith(\"--\"))\n"
        "    .join(\"\\n\");\n"
        "\n"
        "const splitStatements = (sql: string): string[] =>\n"
        "  stripComments(sql)\n"
        "    .split(/;\\s*$/m)\n"
        "    .map((s) => s.trim())\n"
        "    .filter((s) => s.length > 0);",
    )


def section_vibe(doc: Document) -> None:
    add_heading(doc, "8. 바이브코딩 프로세스", level=1)

    add_heading(doc, "8.1 Interface Hub에서 상속받은 CLAUDE.md", level=2)
    add_paragraph(
        doc,
        "Cost Compass 저장소의 CLAUDE.md에는 Interface Hub 개발 중 겪은 이슈를 선제 반영했다. "
        "같은 시행착오를 두 번 겪지 않도록 '하지 말아야 할 것'과 '반드시 해야 할 것'을 명문화한 것이다. "
        "두 번째 프로젝트이기 때문에 반나절에 완성 가능한 속도가 이 문서에서 나왔다.",
    )
    add_bullets(
        doc,
        [
            "Prisma Client import 경로: @/lib/generated/prisma/client (⚠️ NOT @prisma/client).",
            "Prisma 7 모델 타입명: DivisionModel / ProjectModel / CostItemModel (접미사 Model이 붙는다).",
            "schema.prisma에는 datasource.url을 쓰지 않는다 (Prisma 7에서 deprecated).",
            "Recharts ResponsiveContainer에는 minWidth={0} minHeight={0} 필수.",
            "shadcn Select + RHF은 Controller로 감쌀 것 (register 금지).",
            "Server Component의 params는 Promise 타입: const { id } = await params.",
            "금액은 정수(원 단위)로 저장, 소수점 없음. Intl.NumberFormat('ko-KR')로 표시.",
            "shadcn Button은 asChild 미지원 → buttonVariants() + Link 조합.",
            "em dash 문자와 한글 모음(유니코드 U+3161) 사용 금지 (README와 문서 전체). 하이픈만 사용.",
        ],
    )

    add_heading(doc, "8.2 재발 이슈 0건, 재발 방지 이슈 9건", level=2)
    add_paragraph(
        doc,
        "Interface Hub에서 한 번 겪고 CLAUDE.md에 박아둔 9가지 이슈는 Cost Compass 개발 중 "
        "단 한 번도 재발하지 않았다. 프롬프트 상단에 '환경적 제약'으로 명시하는 원칙 덕분에 "
        "AI가 구버전 API나 오답 import로 빗겨가는 일이 사라졌다.",
    )

    add_heading(doc, "8.3 프롬프트 예시", level=2)

    add_paragraph(doc, "예시 1: 본부 드릴다운 추가 프롬프트 발췌", bold=True, space_after=2)
    add_code_block(
        doc,
        "lib/actions/cost-query.ts에 함수 2개 추가.\n"
        "\n"
        "export async function listDivisions(): Promise<Array<{\n"
        "  id: string; name: string; headcount: number;\n"
        "  projectCount: number; activeProjectCount: number;\n"
        "  actual: number; standard: number; varianceRatio: number;\n"
        "}>>\n"
        "\n"
        "export async function getDivisionDetail(id: string): Promise<{\n"
        "  division: DivisionModel;\n"
        "  kpis: { ... };\n"
        "  projects: Array<{ ... }>;\n"
        "  monthlyTrend: Array<{ ... }>;\n"
        "  categoryBreakdown: Array<{ ... }>;\n"
        "} | null>\n"
        "\n"
        "⚠️ params는 Promise, const { id } = await params.\n"
        "⚠️ Recharts Tooltip formatter는 (value, name) 받아 Number()/String() 캐스팅.",
    )
    add_paragraph(
        doc,
        "해설: 반환 타입을 먼저 명문화하고 환경 제약을 프롬프트 상단에 고정했다. AI가 파라미터 순서와 에러 형태를 추측하지 않고 계약에 맞춰 구현한다.",
    )

    add_paragraph(doc, "예시 2: Recharts 3 Tooltip 시그니처 대응 프롬프트", bold=True, space_after=2)
    add_code_block(
        doc,
        "Recharts 3의 Tooltip formatter는 ValueType | undefined가 들어온다.\n"
        "타입 annotation을 value: number, name: string으로 하면 컴파일 에러.\n"
        "\n"
        "해결 방향:\n"
        "- (value, name) 무타입으로 받고 내부에서 Number(value), String(name) 캐스팅.\n"
        "- 모든 Recharts Tooltip 컴포넌트에 동일 패턴을 일괄 적용.",
    )
    add_paragraph(
        doc,
        "해설: 현상(컴파일 에러)과 원인, 수정 방향을 함께 제시해 AI가 버그를 추측하는 시간을 아꼈다. 4개 차트 컴포넌트 모두 한 번에 수정.",
    )

    add_paragraph(doc, "예시 3: CSV 내보내기 서버 액션 프롬프트", bold=True, space_after=2)
    add_code_block(
        doc,
        "lib/actions/export.ts 생성.\n"
        "\n"
        "export async function exportProjectCostCsv(projectId: string):\n"
        "  Promise<\n"
        "    | { ok: true; filename: string; content: string }\n"
        "    | { ok: false; error: string }\n"
        "  >\n"
        "\n"
        "요구:\n"
        "- BOM(\\uFEFF) 포함 UTF-8 (엑셀 한글 호환)\n"
        "- 컬럼: 기간, 카테고리, 부담 본부, 출처 본부, 표준, 실제, Variance(%), 내부대체, 비고\n"
        "- 이스케이프: 쉼표/개행/큰따옴표 포함 시 큰따옴표로 감싸고 내부 큰따옴표는 두 번\n"
        "- 파일명: project-{code}-{yyyymmdd-HHmm}.csv",
    )
    add_paragraph(
        doc,
        "해설: ActionResult 판별 유니온으로 성공/실패를 강제했다. UI 쪽 toast 분기가 타입으로 보호된다.",
    )

    add_heading(doc, "8.4 검증 루프 (Phase당 반복)", level=2)
    add_numbered(
        doc,
        [
            "요구 정의: 데이터 모델, 화면, 동작을 문장으로 정리한다.",
            "프롬프트: 타입 시그니처, 파일 경로, 환경 제약을 명시한다.",
            "생성: Claude Code가 파일을 생성하거나 수정한다.",
            "컴파일 검증: npx tsc --noEmit이 0이어야 한다.",
            "실동작 검증: 브라우저에서 실제 사용자 플로우를 수행하고 curl로 200 응답과 주요 문자열을 확인한다.",
            "커밋: Conventional Commits 규칙으로 단일 논리 단위의 커밋 하나를 만든다.",
        ],
    )

    add_heading(doc, "8.5 소요 시간과 산출", level=2)
    add_table(
        doc,
        headers=["항목", "수치"],
        rows=[
            ["총 소요 시간", "약 4시간 (Phase 0 → Phase 3)"],
            ["생성/수정 파일", "약 30개 (TypeScript, Prisma, Python 스크립트 포함)"],
            ["시드 데이터", "5본부, 22프로젝트, 326 원가 항목 (총 282.9억원)"],
            ["재발 이슈", "0건 (Interface Hub에서 학습한 9가지 이슈 모두 선제 방지)"],
        ],
        col_widths_cm=[5.0, 10.8],
    )


def section_validation(doc: Document) -> None:
    add_heading(doc, "9. 검증 전략", level=1)

    add_heading(doc, "9.1 타입 안전", level=2)
    add_bullets(
        doc,
        [
            "TypeScript strict 모드를 전 코드에 적용한다.",
            "any 사용을 금지하고 unknown을 받은 뒤 타입 가드로 좁힌다.",
            "Recharts의 느슨한 타입은 (value, name) 무타입 + 내부 캐스팅 패턴으로 흡수한다.",
        ],
    )

    add_heading(doc, "9.2 런타임 검증", level=2)
    add_bullets(
        doc,
        [
            "Server Action은 never-throw 계약. CSV 내보내기의 결과는 ActionResult 판별 유니온으로 수렴.",
            "Variance 계산은 standard=0 방어. 0으로 나누지 않는다.",
            "getProjectDetail, getDivisionDetail은 대상이 없으면 null 반환 → 페이지에서 notFound()로 위임.",
        ],
    )

    add_heading(doc, "9.3 UI 회귀 점검", level=2)
    add_bullets(
        doc,
        [
            "각 Phase 완료 후 curl로 200 응답과 주요 한글 문자열 존재를 확인한다.",
            "GET / 에서 '원가 대시보드', '전체 실제 원가', '프로젝트 Top 10' 등이 렌더되는지 확인.",
            "GET /projects 에서 PRJ-001 ~ PRJ-022 22개 코드가 전부 렌더되는지 확인.",
            "GET /divisions/[id] 에서 KPI와 소속 프로젝트 섹션이 렌더되는지 확인.",
        ],
    )

    add_heading(doc, "9.4 시드 기반 데모 가능성", level=2)
    add_bullets(
        doc,
        [
            "prisma/seed.ts가 5본부, 22프로젝트, 326 원가 항목을 생성한다.",
            "카테고리별 금액 범위(LABOR 5천만~3억 등)를 설정해 차트가 의미 있게 그려지게 했다.",
            "15% 확률의 내부대체를 주입해 sourceDivisionId가 비어 있지 않은 레코드를 확보한다.",
            "시드는 upsert와 deleteMany 조합으로 멱등성을 갖는다.",
        ],
    )


def section_troubleshooting(doc: Document) -> None:
    add_heading(doc, "10. 트러블슈팅 사례", level=1)

    add_heading(doc, "10.1 Recharts 3 Tooltip formatter 시그니처 변경", level=2)
    add_paragraph(
        doc,
        "증상: Tooltip의 formatter에 (value: number, name: string) 타입을 달면 "
        "TS2322: Type '(value: number, name: string) => [string, string]' is not assignable to "
        "type 'Formatter<ValueType, NameType>'와 함께 컴파일이 깨졌다.",
    )
    add_paragraph(
        doc,
        "원인: Recharts 3부터 Tooltip Formatter의 ValueType이 string | number | undefined로 넓어져 "
        "정확한 number 타입으로 좁히면 바인딩이 실패한다.",
    )
    add_paragraph(
        doc,
        "해결: formatter 콜백을 (value, name) 무타입으로 받고 내부에서 Number(value).toLocaleString(), "
        "String(name) 캐스팅으로 다룬다. 네 개 차트 컴포넌트에 동일 패턴을 일괄 적용했다. "
        "이 이슈는 CLAUDE.md에 추가해 다음 프로젝트에서 재발하지 않게 한다.",
    )

    add_heading(doc, "10.2 shadcn Button의 asChild 미지원", level=2)
    add_paragraph(
        doc,
        "증상: Button asChild 형태로 Link를 감싸려 하자 "
        "TS2322: Property 'asChild' does not exist on type ButtonProps로 실패했다.",
    )
    add_paragraph(
        doc,
        "원인: 이 프로젝트의 shadcn Button 구현은 Base UI 기반이며 Slot asChild 패턴을 채택하지 않았다.",
    )
    add_paragraph(
        doc,
        "해결: buttonVariants({ variant, size }) 클래스만 뽑아 next/link의 Link에 className으로 적용했다. "
        "동일 비주얼을 유지하면서 Link의 기본 시맨틱(prefetch, 키보드 포커스)을 보존한다.",
    )

    add_heading(doc, "10.3 Prisma 7 모델 타입 접미사", level=2)
    add_paragraph(
        doc,
        "증상: @/lib/generated/prisma/models에서 Division, Project, CostItem을 import하면 "
        "Module has no exported member 'Division'으로 실패했다.",
    )
    add_paragraph(
        doc,
        "원인: Prisma 7은 모델 런타임 타입을 DivisionModel, ProjectModel, CostItemModel로 export한다. "
        "기존 prisma-client-js의 네이밍과 다르다.",
    )
    add_paragraph(
        doc,
        "해결: import를 DivisionModel 등 Model 접미사 버전으로 교체했다. 이 이슈는 Interface Hub에서 "
        "이미 한 번 겪고 CLAUDE.md에 박아뒀지만, Cost Compass 첫 Prisma 셋업 코드에서 초기 문구로 "
        "다시 튀어나와 빠르게 수정했다. 이후 재발 0건.",
    )


def section_appendix(doc: Document) -> None:
    add_heading(doc, "11. 부록", level=1)

    add_heading(doc, "11.1 저장소와 데모", level=2)
    add_bullets(
        doc,
        [
            "저장소: https://github.com/dbwls99706/cost-compass",
            "라이브 데모: https://cost-compass.vercel.app",
            "자매 프로젝트 Interface Hub: https://github.com/dbwls99706/interface-hub",
        ],
    )

    add_heading(doc, "11.2 주요 디렉터리 트리", level=2)
    add_code_block(
        doc,
        "app/\n"
        "  page.tsx                   대시보드 (KPI, 차트 5종, 최근 항목)\n"
        "  projects/\n"
        "    page.tsx                 프로젝트 목록\n"
        "    [id]/page.tsx            프로젝트 상세\n"
        "    [id]/not-found.tsx\n"
        "  divisions/\n"
        "    page.tsx                 본부 목록\n"
        "    [id]/page.tsx            본부 상세\n"
        "    [id]/not-found.tsx\n"
        "  layout.tsx\n"
        "components/\n"
        "  ui/                        shadcn/ui 기반 공통 컴포넌트\n"
        "  dashboard/                 KpiCard, FilterBar, 차트 4종, Badge 3종, ExportCsvButton\n"
        "lib/\n"
        "  actions/\n"
        "    cost-query.ts            대시보드/목록/상세 집계 Server Actions\n"
        "    export.ts                CSV 내보내기 Server Action\n"
        "  prisma.ts                  PrismaLibSql adapter + 환경 자동 감지\n"
        "  types/db.ts                enum, 라벨, 포맷터, Variance 계산\n"
        "  generated/prisma/          Prisma Client (빌드 시 생성)\n"
        "prisma/\n"
        "  schema.prisma              Division, Project, CostItem\n"
        "  seed.ts                    5본부, 22프로젝트, 326 원가 항목\n"
        "  migrations/\n"
        "scripts/\n"
        "  turso-deploy.ts            Turso 마이그레이션 적용\n"
        "  turso-reset.ts             _applied_migrations 초기화\n"
        "docs/\n"
        "  screenshots/               README와 문서용 스크린샷 5장\n"
        "  build_planning_doc.py      기획서 생성 스크립트\n"
        "  build_dev_doc.py           개발문서 생성 스크립트\n"
        "  기획서.docx, 개발문서.docx   제출용 문서",
    )

    add_heading(doc, "11.3 package.json scripts", level=2)
    add_code_block(
        doc,
        "\"scripts\": {\n"
        "  \"dev\": \"next dev\",\n"
        "  \"prebuild\": \"prisma generate\",\n"
        "  \"build\": \"next build\",\n"
        "  \"start\": \"next start\",\n"
        "  \"lint\": \"eslint\",\n"
        "  \"postinstall\": \"prisma generate\",\n"
        "  \"db:deploy\": \"tsx scripts/turso-deploy.ts\",\n"
        "  \"db:reset:remote\": \"tsx scripts/turso-reset.ts\",\n"
        "  \"db:seed:remote\": \"tsx prisma/seed.ts\"\n"
        "}",
    )

    add_heading(doc, "11.4 주요 환경변수", level=2)
    add_table(
        doc,
        headers=["변수", "용도", "예시"],
        rows=[
            ["DATABASE_URL", "로컬 SQLite 또는 Turso 단일 URL", "file:./dev.db"],
            ["TURSO_DATABASE_URL", "프로덕션(Turso) URL", "libsql://project.turso.io"],
            ["TURSO_AUTH_TOKEN", "Turso 인증 토큰", "eyJhbGciOi..."],
        ],
        col_widths_cm=[4.2, 5.6, 6.0],
    )

    add_heading(doc, "11.5 자매 프로젝트 크로스 레퍼런스", level=2)
    add_table(
        doc,
        headers=["공통 자산", "Interface Hub", "Cost Compass"],
        rows=[
            ["lib/prisma.ts", "libSQL adapter + 환경 자동 감지", "동일 구조 복제"],
            ["scripts/turso-deploy.ts", "마이그레이션 파서 (라인 단위 주석 제거)", "동일 구조 복제"],
            ["scripts/turso-reset.ts", "_applied_migrations 초기화", "동일 구조 복제"],
            ["CLAUDE.md", "9가지 선제 이슈 정의", "동일 이슈 상속 + 본 프로젝트 전용 추가"],
            ["docs/build_*.py", "기획서/개발문서 생성 스크립트", "동일 구조 복제 + Cost Compass 내용"],
        ],
        col_widths_cm=[4.2, 5.8, 6.0],
    )


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = KOR_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)

    build_cover(doc)
    build_toc(doc)
    section_scope(doc)
    section_architecture(doc)
    section_data_model(doc)
    section_aggregation(doc)
    section_frontend(doc)
    section_drilldown(doc)
    section_turso(doc)
    section_vibe(doc)
    section_validation(doc)
    section_troubleshooting(doc)
    section_appendix(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
