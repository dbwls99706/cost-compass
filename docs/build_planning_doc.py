"""docs/기획서.docx 생성 스크립트."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
SCREENSHOTS = DOCS_DIR / "screenshots"
OUTPUT = DOCS_DIR / "기획서.docx"

KOR_FONT = "맑은 고딕"


def set_korean_font(run, size_pt: float, bold: bool = False, color=None) -> None:
    run.font.name = KOR_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    align=None,
    space_after: float = 6,
    color=None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_korean_font(run, size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    if level == 0:
        size = 22
    elif level == 1:
        size = 16
    elif level == 2:
        size = 13
    else:
        size = 12
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_korean_font(run, size, bold=True, color=RGBColor(0x1F, 0x2D, 0x3D))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        if p.runs:
            run = p.runs[0]
            run.text = item
        else:
            run = p.add_run(item)
        set_korean_font(run, 11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_korean_font(run, 9.5, color=RGBColor(0x55, 0x65, 0x77))
    run.italic = True


def add_image(doc: Document, path: Path, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_korean_font(run, 10.5, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_korean_font(run, 10.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    add_paragraph(
        doc,
        "Cost Compass",
        size=28,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x10, 0x2A, 0x43),
        space_after=4,
    )
    add_paragraph(
        doc,
        "본부/프로젝트 원가 분석 대시보드 기획서",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_paragraph(
        doc,
        "노아에이티에스 2025년 상반기 연구소 인력 채용 포트폴리오 (주제 2)",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x44, 0x55, 0x66),
        space_after=80,
    )

    add_paragraph(
        doc,
        "작성자: 홍유진 (GitHub: dbwls99706)",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "작성일: 2026-04-20",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "라이브 데모: https://cost-compass.vercel.app",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "저장소: https://github.com/dbwls99706/cost-compass",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "자매 프로젝트 (주제 1): https://github.com/dbwls99706/interface-hub",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_heading(doc, "목차", level=0)
    items = [
        "1. 프로젝트 개요",
        "2. 문제 정의",
        "3. 타깃 사용자",
        "4. MVP 범위와 경계",
        "5. 핵심 기능 정의",
        "6. 성공 지표",
        "7. 확장 로드맵",
        "8. 부록",
    ]
    for item in items:
        add_paragraph(doc, item, size=12, space_after=4)
    doc.add_page_break()


def section_overview(doc: Document) -> None:
    add_heading(doc, "1. 프로젝트 개요", level=1)
    add_paragraph(
        doc,
        "Cost Compass는 금융사의 5개 본부와 22개 프로젝트에 걸쳐 분산된 원가 데이터를 "
        "단일 화면에서 집계, 비교, 드릴다운하는 관리회계 대시보드다. 본부장, 관리회계 담당자, "
        "프로젝트 PM이 공통으로 사용할 수 있는 조회 전용 도구로 설계했다. "
        "자매 프로젝트인 Interface Hub가 운영자의 '쓰는' 도구였다면, Cost Compass는 관리자의 '읽는' 도구다.",
    )
    add_heading(doc, "1.1 프로젝트 메타", level=2)
    add_table(
        doc,
        headers=["항목", "내용"],
        rows=[
            ["프로젝트명", "Cost Compass"],
            ["라이브 데모", "https://cost-compass.vercel.app"],
            ["저장소", "https://github.com/dbwls99706/cost-compass"],
            ["자매 프로젝트", "Interface Hub (주제 1, https://github.com/dbwls99706/interface-hub)"],
            ["개발 기간", "반나절 (Claude Code 기반 바이브코딩)"],
            ["기술 스택 요약", "Next.js 16, Prisma 7, Turso libSQL, Tailwind, Recharts, Vercel"],
            ["데이터 규모", "5본부, 22프로젝트, 326개 원가 항목 (시드)"],
        ],
        col_widths_cm=[3.8, 12.0],
    )


def section_problem(doc: Document) -> None:
    add_heading(doc, "2. 문제 정의", level=1)

    add_heading(doc, "2.1 배경", level=2)
    add_paragraph(
        doc,
        "은행 등 금융사는 5개 안팎의 본부 조직이 동시에 20여 개의 프로젝트를 수행한다. "
        "각 프로젝트는 인건비, 외주비, 운영비, 공통배부 등 여러 카테고리의 원가를 월 단위로 집행하고, "
        "본부 간 내부대체(inter-unit transfer)도 수시로 발생한다. "
        "원가 데이터가 본부 시스템마다 분산되면 관리자 관점의 통합 시야를 확보하기 어렵다.",
    )

    add_heading(doc, "2.2 현재의 문제점", level=2)
    add_bullets(
        doc,
        [
            "본부 시스템마다 원가 관리 도구가 흩어져 있어 전사 총액과 본부별 비중을 한눈에 보기 어렵다.",
            "내부대체가 섞이면 실제 부담 본부와 출처 본부가 얽혀 원가 책임 구조가 흐려진다.",
            "표준원가와 실제원가의 Variance를 월/카테고리 축으로 드릴다운하려면 엑셀을 매번 재가공해야 한다.",
            "관리자 관점의 전사 KPI(총 실제 원가, 진행 프로젝트, 내부대체 비중 등)가 한 화면에 없다.",
        ],
    )

    add_heading(doc, "2.3 해결 방향", level=2)
    add_bullets(
        doc,
        [
            "본부, 프로젝트, 원가 항목을 단일 데이터 모델로 정규화한다 (Division, Project, CostItem 3개 엔티티).",
            "본부, 프로젝트, 카테고리, 월의 4개 축으로 드릴다운이 가능한 조회 전용 대시보드를 제공한다.",
            "URL 쿼리스트링에 필터 상태를 저장해 링크 하나로 동일한 뷰를 공유할 수 있게 한다.",
            "차트의 막대 클릭은 그대로 해당 본부/프로젝트 상세로 이동해 분석 흐름을 단절 없이 잇는다.",
        ],
    )


def section_users(doc: Document) -> None:
    add_heading(doc, "3. 타깃 사용자", level=1)
    add_paragraph(
        doc,
        "Cost Compass는 관리 계층과 실무 계층의 사용자를 동시에 고려해 설계했다. "
        "역할에 따라 보고 싶은 집계의 결이 다르지만 같은 데이터 모델 위에서 필터와 드릴다운으로 해결한다.",
    )
    add_table(
        doc,
        headers=["역할", "주요 페인포인트", "Cost Compass가 제공하는 가치"],
        rows=[
            [
                "본부장 (1차)",
                "자기 본부의 실제 원가와 표준 대비 Variance가 어디서 발생하는지 월말이 돼야 알 수 있다.",
                "본부 상세 화면에서 KPI, 월별 추이, 카테고리 비중, 소속 프로젝트 전 건을 한 페이지에서 확인한다.",
            ],
            [
                "관리회계 담당자 (2차)",
                "전사 원가 집계와 프로젝트별 Variance 분석을 매달 엑셀로 재가공한다.",
                "대시보드에서 본부, 프로젝트, 카테고리 축으로 즉시 드릴다운하고 CSV로 내보낸다.",
            ],
            [
                "프로젝트 PM (3차)",
                "자기 프로젝트의 원가 집행 현황을 본부 시스템에서 별도로 받아야 한다.",
                "프로젝트 상세에서 월별 집행, 카테고리 비중, 내부대체 출처까지 즉시 확인한다.",
            ],
        ],
        col_widths_cm=[3.5, 6.0, 6.0],
    )


def section_mvp(doc: Document) -> None:
    add_heading(doc, "4. MVP 범위와 경계", level=1)
    add_paragraph(
        doc,
        "반나절이라는 개발 기간 안에 '읽는' 도구로서의 핵심 가치를 입증하기 위해 범위를 명확히 잘랐다. "
        "조회와 드릴다운, 필터 공유에 집중하고 데이터 생성/수정은 범위에서 제외했다.",
    )

    add_heading(doc, "4.1 MoSCoW 분석", level=2)
    add_table(
        doc,
        headers=["우선순위", "범위"],
        rows=[
            ["Must", "대시보드 KPI 4종, 차트 5종, 프로젝트 목록/상세, 본부 드릴다운, CSV 내보내기"],
            ["Should", "URL 기반 필터 공유, Variance 30% 초과 강조, 차트 클릭 드릴다운"],
            ["Could", "표준원가 배분 규칙 엔진(ABC), 월 마감 확정 워크플로"],
            ["Won't", "원가 직접 입력 폼, 권한/인증(RLS), 실ERP 연동, Excel 정식 템플릿"],
        ],
        col_widths_cm=[3.2, 12.6],
    )

    add_heading(doc, "4.2 왜 조회 전용인가", level=2)
    add_bullets(
        doc,
        [
            "자매 프로젝트 Interface Hub는 CRUD와 실행을 포함한 '쓰는' 도구였다. Cost Compass는 대비되는 '읽는' 도구로 포지셔닝한다.",
            "원가 입력 워크플로는 ERP 영역에 속한다. 대시보드는 통합 조회에 집중하고, 입력은 확장 로드맵에 남긴다.",
            "시드 데이터 326건으로 차트와 집계 로직의 의미 있는 분포를 실증한다 (총 실제 원가 약 282.9억원).",
        ],
    )


def section_features(doc: Document) -> None:
    add_heading(doc, "5. 핵심 기능 정의", level=1)
    add_paragraph(
        doc,
        "MVP가 제공하는 기능은 크게 5가지 화면으로 구성된다. 각 화면은 시나리오에 맞춘 데이터 집계와 인터랙션을 함께 갖춘다.",
    )

    add_heading(doc, "5.1 전사 대시보드", level=2)
    add_paragraph(
        doc,
        "전체 실제 원가, 진행 프로젝트 수, Variance 비율, 내부대체 비중의 KPI 4종을 상단에 노출한다. "
        "월별 추이 LineChart, 카테고리 비중 PieChart, 본부별 카테고리 Stacked BarChart, "
        "프로젝트 Top 10 Horizontal BarChart, 최근 집행 원가 10건 테이블까지 한 페이지에 배치한다. "
        "상단 FilterBar의 기간/본부/상태 필터는 URL 쿼리스트링에 반영되어 링크로 그대로 공유된다.",
    )
    add_image(doc, SCREENSHOTS / "01-dashboard.png")
    add_caption(doc, "전사 KPI, 월별 추이, 카테고리 비중, 본부별 스택, 프로젝트 Top 10을 한 화면에 집계한다.")

    add_heading(doc, "5.2 프로젝트 목록", level=2)
    add_paragraph(
        doc,
        "22개 프로젝트를 표 형식으로 나열하고 본부, 상태로 필터링한다. Variance가 30%를 초과한 행은 "
        "배경색으로 강조해 한 번의 시선으로 위험 프로젝트를 식별한다. 행 클릭은 그대로 프로젝트 상세로 이동한다.",
    )
    add_image(doc, SCREENSHOTS / "02-project-list.png")
    add_caption(doc, "본부와 상태로 필터링하고 Variance 30% 초과 프로젝트는 행 배경으로 강조한다.")

    add_heading(doc, "5.3 프로젝트 상세", level=2)
    add_paragraph(
        doc,
        "표준원가, 실제원가, Variance KPI 3종 위에 카테고리 비중 PieChart와 월별 추이 LineChart를 배치하고, "
        "하단에 원가 항목 상세 테이블을 둔다. 항목마다 기간, 카테고리, 부담 본부, 내부대체 출처 본부를 "
        "모두 노출해 금액이 어디서 발생했고 어떻게 이동했는지 추적할 수 있다. 우상단 'CSV 내보내기' 버튼은 "
        "UTF-8 BOM 포함 CSV를 생성해 엑셀에서 한글이 깨지지 않도록 한다.",
    )
    add_image(doc, SCREENSHOTS / "03-project-detail.png")
    add_caption(
        doc,
        "표준원가, 실제원가, Variance KPI와 카테고리 비중, 월별 추이, 원가 항목 상세를 드릴다운한다. 우상단에서 UTF-8 BOM CSV로 내보낸다.",
    )

    add_heading(doc, "5.4 본부 목록", level=2)
    add_paragraph(
        doc,
        "5개 본부를 카드 그리드로 나열하고 실제 원가 내림차순으로 정렬한다. 각 카드는 인원수, "
        "프로젝트 수와 진행 프로젝트 수, 실제 원가, 표준 대비 Variance를 동시에 보여준다. "
        "카드 전체가 링크라 본부 상세로 자연스럽게 이어진다. 대시보드의 본부별 Stacked Bar를 클릭해도 동일한 상세 화면으로 이동한다.",
    )
    add_image(doc, SCREENSHOTS / "04-divisions.png")
    add_caption(doc, "5개 본부를 실제 원가 내림차순으로 정렬해 비교한다.")

    add_heading(doc, "5.5 본부 상세", level=2)
    add_paragraph(
        doc,
        "본부 단위 KPI 4종(실제 원가, 표준 원가, 진행/전체 프로젝트 수, 내부대체 비중) 위에 월별 추이와 "
        "카테고리 비중 차트를 배치하고, 그 아래 소속 프로젝트 전 건을 실제 원가 내림차순으로 나열한다. "
        "대시보드의 전체 시야에서 본부 시야로, 다시 프로젝트 시야로 단절 없이 드릴다운할 수 있다.",
    )
    add_image(doc, SCREENSHOTS / "05-division-detail.png")
    add_caption(doc, "본부 단위 KPI와 월별 추이, 카테고리 비중, 소속 프로젝트 목록까지 한 번에 확인한다.")


def section_metrics(doc: Document) -> None:
    add_heading(doc, "6. 성공 지표", level=1)

    add_heading(doc, "6.1 MVP 검증 기준", level=2)
    add_bullets(
        doc,
        [
            "시드 326건이 대시보드 KPI에 정확히 집계되고 전체 실제 원가 합계가 약 282.9억원으로 일치한다.",
            "본부별 Stacked BarChart의 막대 클릭이 해당 본부 상세로 이동하고, 프로젝트 Top 10의 막대 클릭이 프로젝트 상세로 이동한다.",
            "CSV 내보내기 파일이 엑셀에서 한글 깨짐 없이 열리고 BOM 포함 UTF-8로 저장된다.",
            "로컬 SQLite와 Turso libSQL을 환경변수만으로 전환할 수 있고, 코드 수정이 없다.",
            "TypeScript strict 기준 npx tsc --noEmit 0건과 npm run build 성공을 통과한다.",
        ],
    )

    add_heading(doc, "6.2 운영 KPI (제품이 운영에 투입된 가정)", level=2)
    add_table(
        doc,
        headers=["KPI", "정의", "기대 효과"],
        rows=[
            [
                "월별 예산 대비 실적 편차",
                "본부/프로젝트별 실제와 표준의 Variance 추이",
                "예산 수립 정밀도 개선",
            ],
            [
                "Variance 30% 초과 프로젝트 수",
                "표준 대비 실제가 30% 이상 차이 난 프로젝트 건수",
                "위험 프로젝트 조기 식별",
            ],
            [
                "내부대체 비중",
                "전체 실제 원가 대비 내부대체 금액 비율",
                "본부 간 원가 이동 투명성 확보",
            ],
            [
                "카테고리별 원가 구조",
                "인건비/외주비/운영비/공통배부의 비중 변화",
                "원가 체계 재설계 근거 확보",
            ],
        ],
        col_widths_cm=[3.6, 5.8, 6.4],
    )


def section_roadmap(doc: Document) -> None:
    add_heading(doc, "7. 확장 로드맵", level=1)
    add_paragraph(
        doc,
        "MVP는 반나절 기준 산출물이며 실제 관리회계 도구로 확장할 경로를 다음과 같이 설계한다. "
        "각 Phase는 이전 Phase의 데이터 모델을 그대로 활용한다.",
    )
    add_table(
        doc,
        headers=["Phase", "범위", "예상 공수", "선행 의존"],
        rows=[
            ["Phase 2 (M+1)", "원가 직접 입력 폼과 월 마감/확정 워크플로", "약 2주", "없음"],
            ["Phase 3 (M+2)", "표준원가 배분 규칙 엔진 (ABC 기반)", "약 2주", "Phase 2"],
            ["Phase 4 (M+3)", "본부 단위 권한 (Row Level Security)", "약 1주", "Phase 2"],
            ["Phase 5 (M+4)", "Excel 정식 템플릿 내보내기 (xlsx, 셀 서식 포함)", "약 3일", "Phase 2"],
            ["Phase 6 (M+5)", "외부 ERP 연동 (정산 잔액 동기화)", "약 2주", "Phase 3"],
        ],
        col_widths_cm=[3.0, 6.5, 2.6, 3.6],
    )


def section_appendix(doc: Document) -> None:
    add_heading(doc, "8. 부록", level=1)

    add_heading(doc, "8.1 용어집", level=2)
    add_table(
        doc,
        headers=["용어", "설명"],
        rows=[
            ["표준원가", "사전에 수립한 계획 단가와 수량을 기준으로 산출한 원가."],
            ["실제원가", "실제로 집행된 비용. 표준원가와의 차이를 Variance로 부른다."],
            ["Variance", "실제와 표준의 차이. (실제 - 표준) 금액과 표준 대비 비율 두 축으로 본다."],
            ["내부대체", "한 본부가 부담한 비용을 다른 본부 원천에서 이전받는 회계 처리. sourceDivisionId로 출처 본부를 기록한다."],
            ["ABC", "Activity-Based Costing. 활동 단위의 원가 동인을 기준으로 간접비를 배부하는 방식."],
            ["RLS", "Row Level Security. 행 단위 접근 제어. 본부별로 데이터 가시성을 분리한다."],
        ],
        col_widths_cm=[3.2, 12.6],
    )

    add_heading(doc, "8.2 자매 프로젝트 Interface Hub와의 비교", level=2)
    add_table(
        doc,
        headers=["항목", "Interface Hub", "Cost Compass"],
        rows=[
            ["성격", "운영자 도구 (CRUD + 실행)", "분석 도구 (조회 + 드릴다운)"],
            ["주제", "인터페이스 통합 관리", "원가/관리회계"],
            ["핵심 패턴", "Adapter 패턴, SWR 조건부 폴링", "JS 집계, Server Component 위주"],
            ["차트", "1종 (BarChart)", "5종 (Line, Pie, Stacked, Horizontal Bar, Table)"],
            ["개발 기간", "1일", "반나절"],
        ],
        col_widths_cm=[3.2, 6.2, 6.4],
    )

    add_heading(doc, "8.3 참고 자료", level=2)
    add_bullets(
        doc,
        [
            "노아에이티에스 2025년 상반기 연구소 인력 채용 공고 (주제 2 관리회계 대시보드)",
            "자매 프로젝트 Interface Hub 저장소: https://github.com/dbwls99706/interface-hub",
            "Prisma 공식 문서: libSQL Driver Adapter",
            "Turso 공식 문서: Edge SQLite 서비스",
            "Next.js 16 공식 문서: App Router, Server Actions",
            "Recharts 공식 문서: ResponsiveContainer, LineChart, PieChart, BarChart",
        ],
    )


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = KOR_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), KOR_FONT)
    r_fonts.set(qn("w:ascii"), KOR_FONT)
    r_fonts.set(qn("w:hAnsi"), KOR_FONT)

    build_cover(doc)
    build_toc(doc)
    section_overview(doc)
    section_problem(doc)
    section_users(doc)
    section_mvp(doc)
    section_features(doc)
    section_metrics(doc)
    section_roadmap(doc)
    section_appendix(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
